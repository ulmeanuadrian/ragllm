# rag_api.py - VERSIUNEA PENTRU CRUD

import os
import re
import time
import pypdf
import docx
import json
import traceback
import logging
from datetime import datetime
from pathlib import Path
from fastapi import FastAPI, Request, UploadFile, File, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, Response
from qdrant_client import QdrantClient
from typing import List, Dict, Any, Optional
from models import CollectionCreate, DocumentDeleteRequest, GenerateRequest
from pydantic import ConfigDict, BaseModel, Field
from typing import Dict, List, Optional, Any

# Implementăm o versiune simplificată a clasei Document pentru a evita problemele cu pydantic
class Document(BaseModel):
    model_config = ConfigDict(arbitrary_types_allowed=True)
    content: str
    content_type: str = "text"
    id: Optional[str] = None
    score: Optional[float] = None
    meta: Dict[str, Any] = Field(default_factory=dict)
    embedding: Optional[List[float]] = None
    id_hash_keys: Optional[List[str]] = None
# Importuri pentru farm-haystack 1.26.4
from haystack.preview.components.preprocessors import PreProcessor
from haystack.document_stores.qdrant import QdrantDocumentStore
from haystack.preview.components.retrievers import EmbeddingRetriever
from haystack.preview.pipelines import Pipeline

# Configurare logging structurat în format JSON
class JSONFormatter(logging.Formatter):
    def format(self, record):
        # Creăm un dicționar de bază cu informații standard
        log_record = {
            "timestamp": self.formatTime(record, datefmt="%Y-%m-%dT%H:%M:%S"),
            "level": record.levelname,
            "message": record.getMessage(),
            "module": record.module,
            "function": record.funcName,
            "line": record.lineno
        }
        
        # Adăugăm toate atributele extra din LogRecord
        for key, value in record.__dict__.items():
            # Excludem atributele standard și interne
            if key not in ["args", "exc_info", "exc_text", "msg", "message", "levelname", 
                          "levelno", "pathname", "filename", "module", "lineno", 
                          "funcName", "created", "asctime", "msecs", "relativeCreated",
                          "thread", "threadName", "processName", "process", "name"]:
                log_record[key] = value
                
        return json.dumps(log_record)

# Configurare logger
logger = logging.getLogger("rag_api")
handler = logging.StreamHandler()
handler.setFormatter(JSONFormatter())
logger.addHandler(handler)
logger.setLevel(logging.INFO)

# Inițializare modele de embedding și alte componente la nivel global
doc_embedder = SentenceTransformersDocumentEmbedder(
    model="sentence-transformers/all-mpnet-base-v2",
    batch_size=32
)

text_embedder = SentenceTransformersTextEmbedder(
    model="sentence-transformers/all-mpnet-base-v2",
    batch_size=32
)

# Document splitter pentru fragmentarea documentelor
doc_splitter = DocumentSplitter(
    split_by="word",
    split_length=100,
    split_overlap=20
)

# Helper pentru măsurarea duratei operațiilor și logging structurat
class TimingLogger:
    def __init__(self, operation_name, collection=None, **extra_fields):
        self.operation_name = operation_name
        self.collection = collection
        self.start_time = None
        self.extra_fields = extra_fields
        
    def __enter__(self):
        self.start_time = time.time()
        logger.info(f"Începere {self.operation_name}", 
                   extra={"collection": self.collection, **self.extra_fields})
        return self
        
    def __exit__(self, exc_type, exc_val, exc_tb):
        duration = time.time() - self.start_time
        extra = {
            "duration": f"{duration:.2f}s",
            "duration_ms": int(duration * 1000),
            "timestamp_end": datetime.now().isoformat()
        }
        
        if self.collection:
            extra["collection"] = self.collection
            
        # Adăugăm orice câmpuri extra specificate la inițializare
        extra.update(self.extra_fields)
            
        if exc_type is not None:
            # Dacă a apărut o excepție, o logăm ca eroare cu detalii complete
            error_details = traceback.format_exception(exc_type, exc_val, exc_tb)
            logger.error(f"{self.operation_name} eșuat în {duration:.2f}s: {exc_val}", 
                        extra={**extra, "error": str(exc_val), "traceback": error_details})
        else:
            # Altfel, logăm ca succes cu metrici de performanță
            logger.info(f"{self.operation_name} finalizat în {duration:.2f}s", extra=extra)

# Import pentru generatorul Gemini

# Importăm generatorul Gemini și utilitățile
from gemini_generator import GeminiGenerator
from utils import process_pdf_optimized, process_json_chunks

# Cache pentru interogări pentru a evita procesarea repetată a acelorași întrebări
from collections import OrderedDict
QUERY_CACHE = OrderedDict()
MAX_QUERY_CACHE_SIZE = 100  # Limităm dimensiunea cache-ului pentru a evita consumul excesiv de memorie

def get_from_query_cache(key: str):
    """Obține rezultatele din cache pentru o interogare."""
    global QUERY_CACHE
    if key in QUERY_CACHE:
        # Mutăm intrarea la sfârșitul OrderedDict pentru a o marca ca fiind recent utilizată
        value = QUERY_CACHE.pop(key)
        QUERY_CACHE[key] = value
        logger.info("Rezultate obținute din cache", extra={"cache_key": key[:30]})
        return value
    return None

def save_to_query_cache(key: str, value: list):
    """Salvează rezultatele unei interogări în cache."""
    global QUERY_CACHE, MAX_QUERY_CACHE_SIZE
    
    # Dacă cache-ul a atins dimensiunea maximă, eliminăm cea mai veche intrare
    if len(QUERY_CACHE) >= MAX_QUERY_CACHE_SIZE:
        QUERY_CACHE.popitem(last=False)  # Eliminăm primul element (cel mai vechi)
    
    # Adăugăm noile rezultate în cache
    QUERY_CACHE[key] = value
    logger.info("Rezultate salvate în cache", extra={"cache_key": key[:30], "results_count": len(value)})

# --- CONFIGURARE GENERALĂ ---
UPLOAD_DIR = Path("./uploaded_files")
UPLOAD_DIR.mkdir(exist_ok=True)

qdrant_client = QdrantClient(url="http://localhost:6333")

app = FastAPI(title="RAG API - CRUD")

# Eveniment de startup pentru inițializarea aplicației
@app.on_event("startup")
def startup_event():
    logger.info("Inițializare aplicație")
    try:
        # Verificăm conexiunea la Qdrant
        collections_response = qdrant_client.get_collections()
        logger.info(f"Conexiune la Qdrant stabilită cu succes. Colecții disponibile: {len(collections_response.collections)}")
    except Exception as e:
        logger.error(f"Eroare la conectarea la Qdrant: {str(e)}")
        # Continuăm totuși execuția

# Configurare CORS - cu origini specifice pentru a evita erorile CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:8080", "http://127.0.0.1:8080"],  # Origini specifice
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],  # Metode specifice
    allow_headers=["*"],  # Permite orice header
    expose_headers=["Content-Type", "X-Content-Type-Options"],
    max_age=600  # Cache pentru preflight requests (10 minute)
)

# Middleware pentru gestionarea erorilor
@app.middleware("http")
async def error_handling_middleware(request: Request, call_next):
    try:
        return await call_next(request)
    except Exception as e:
        # Logare eroare server
        print(f"Eroare server: {str(e)}")
        
        # Returnare răspuns de eroare formatat
        return JSONResponse(
            status_code=500,
            content={
                "status": "error",
                "message": "A apărut o eroare la procesarea cererii.",
                "error_type": type(e).__name__
            }
        )

# --- COMPONENTE HAYSTACK REUTILIZABILE ---
# Folosim un model mai performant pentru embeddings (all-mpnet-base-v2 are performanțe mai bune decât all-MiniLM-L6-v2)
# În farm-haystack 1.15.0, nu avem nevoie să inițializăm modelele global, vom folosi EmbeddingRetriever direct în funcții

# În farm-haystack 1.26.4, folosim PreProcessor pentru fragmentarea documentelor
# Inițializăm preprocessor-ul pentru fragmentare - versiune compatibilă cu farm-haystack 1.26.4
preprocessor = PreProcessor(
    clean_empty_lines=True,
    clean_whitespace=True,
    clean_header_footer=True,
    split_by="word",
    split_length=300,
    split_overlap=50,
    split_respect_sentence_boundary=True
)

# --- FUNCȚII HELPER ---
# Dimensiunea vectorului de embedding pentru modelul all-mpnet-base-v2
EMBEDDING_DIM = 768  # Modelul all-mpnet-base-v2 are dimensiunea 768

def get_document_store(collection_name: str, recreate: bool = False) -> QdrantDocumentStore:
    """
    Obține un document store pentru o colecție specificată.
    Gestionează automat recrearea colecției dacă dimensiunea vectorului nu se potrivește.
    
    Args:
        collection_name: Numele colecției
        recreate: Dacă trebuie recreată colecția
    
    Returns:
        QdrantDocumentStore: Document store-ul pentru colecția specificată
    """
    try:
        # Verificăm dacă colecția există deja
        collections_response = qdrant_client.get_collections()
        existing_collections = [c.name for c in collections_response.collections]
        
        # Dacă folderul există și nu dorim să-l recreăm, verificăm dimensiunea vectorului
        if collection_name in existing_collections and not recreate:
            collection_info = qdrant_client.get_collection(collection_name)
            existing_dim = collection_info.config.params.vectors.size
            
            # Dacă dimensiunea nu se potrivește, recreăm folderul
            if existing_dim != EMBEDDING_DIM:
                print(f"Dimensiunea vectorului nu se potrivește: {existing_dim} vs {EMBEDDING_DIM}")
                return get_document_store(collection_name, recreate=True)
        
        # Inițializăm document store cu parametrii optimi
        document_store = QdrantDocumentStore(
            url="http://localhost:6333",
            index=collection_name,
            embedding_dim=EMBEDDING_DIM,
            recreate_index=recreate,
            similarity="cosine"  # Specifiăm explicit metrica de similaritate
        )
        return document_store
    except ValueError as e:
        if "vector size" in str(e) and not recreate:
            logger.warning(f"Recreare folder cu dimensiunea corectă a vectorului: {collection_name}", 
                         extra={"collection": collection_name})
            # Recreăm folderul cu parametrul recreate_index=True
            return get_document_store(collection_name, recreate=True)
        else:
            # Dacă e altă eroare sau am încercat deja recrearea, propagăm eroarea
            logger.error(f"Eroare la inițializarea document store: {str(e)}", 
                       extra={"collection": collection_name, "error": str(e)})
            raise

# --- ENDPOINTS API PENTRU CRUD ---

# CREATE: Endpoint pentru a încărca și procesa fișiere (PDF, DOCX, TXT)
@app.post("/collections/{collection_name}/upload", tags=["CRUD"])
async def upload_file(collection_name: str, file: UploadFile = File(...)):
    """
    Endpoint pentru încărcarea și procesarea fișierelor în foldere.
    Acceptă formate PDF, DOCX, TXT și JSON.
    
    Args:
        collection_name: Numele folder în care se încarcă fișierul
        file: Fișierul încărcat
    
    Returns:
        Informații despre procesarea fișierului
    """
    try:
        start_time = time.time()
        print(f"Procesare fișier: {file.filename} pentru folder: {collection_name}")
        
        # Creăm un nume de fișier unic pentru a evita suprascrierea
        timestamp = int(time.time() * 1000)
        file_path = UPLOAD_DIR / f"{timestamp}_{file.filename}"
        
        # Scriem fișierul pe disc
        with open(file_path, "wb") as f:
            content_bytes = await file.read()
            print(f"Dimensiune fișier: {len(content_bytes)} bytes")
            f.write(content_bytes)

        content = ""
        print(f"Tip fișier detectat: {file.content_type}")
        
        # Procesare în funcție de tipul fișierului
        file_ext = file.filename.lower()
        
        if file_ext.endswith(".pdf"):
            logger.info("Procesare PDF...", extra={"file": file.filename})
            try:
                # Procesare optimizată pentru PDF-uri mari folosind chunk-uri
                content = process_pdf_optimized(file_path)
            except Exception as e:
                logger.error(f"Eroare la procesarea PDF: {str(e)}", extra={"file": file.filename})
                raise HTTPException(status_code=400, detail=f"Eroare la procesarea PDF: {str(e)}")
                
        elif file_ext.endswith(".docx"):
            logger.info("Procesare DOCX...", extra={"file": file.filename})
            try:
                with TimingLogger("Procesare DOCX"):
                    doc = docx.Document(file_path)
                    content_parts = []
                    
                    # Extragem text din paragrafe - folosim list comprehension pentru eficiență
                    content_parts.extend([para.text for para in doc.paragraphs if para.text.strip()])
                    
                    # Extragem text din tabele - optimizăm folosind list comprehension și join
                    for table in doc.tables:
                        content_parts.extend([" | ".join(cell.text for cell in row.cells if cell.text.strip()) 
                                            for row in table.rows if any(cell.text.strip() for cell in row.cells)])
                    
                    # Join o singură dată la final pentru eficiență
                    content = "\n".join(content_parts)
            except Exception as e:
                logger.error(f"Eroare la procesarea DOCX: {str(e)}", extra={"file": file.filename})
                raise HTTPException(status_code=400, detail=f"Eroare la procesarea DOCX: {str(e)}")
                
        elif file_ext.endswith(".txt"):
            logger.info("Procesare TXT...", extra={"file": file.filename})
            try:
                # Folosim un context manager și încercăm mai multe encoding-uri
                encodings = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]
                content = None
                
                for encoding in encodings:
                    try:
                        with open(file_path, "r", encoding=encoding) as f:
                            content = f.read()
                        logger.info(f"Fișierul TXT decodat cu succes folosind encoding: {encoding}", 
                                   extra={"file": file.filename, "encoding": encoding})
                        break
                    except UnicodeDecodeError:
                        continue
                
                if content is None:
                    raise ValueError("Nu s-a putut decoda fișierul cu niciun encoding încercat")
                    
            except Exception as e:
                logger.error(f"Eroare la procesarea TXT: {str(e)}", extra={"file": file.filename})
                raise HTTPException(status_code=400, detail=f"Eroare la procesarea TXT: {str(e)}")
                
        elif file_ext.endswith(".json"):
            logger.info("Procesare JSON optimizată...", extra={"file": file.filename})
            try:
                with TimingLogger("Procesare JSON chunkizat", collection=collection_name, file=file.filename):
                    # Folosim funcția optimizată pentru procesarea fișierelor JSON chunkizate
                    chunks_data = process_json_chunks(file_path)
                
                if chunks_data:
                    logger.info(f"Detectate {len(chunks_data)} chunk-uri predefinite în JSON", 
                               extra={"chunks_count": len(chunks_data), "file": file.filename})
                    
                    # Creăm document store și configurăm pipeline-ul de indexare
                    document_store = get_document_store(collection_name)
                    
                    # Creăm toate documentele într-o singură listă pentru procesare în batch
                    docs_to_index = []
                    
                    for chunk_data in chunks_data:
                        # Adăugăm metadate standard
                        metadata = chunk_data["metadata"]
                        metadata["source"] = file.filename
                        metadata["created_at"] = time.strftime("%Y-%m-%d %H:%M:%S")
                        
                        # Creăm documentul
                        doc = Document(
                            content=chunk_data["content"],
                            meta=metadata
                        )
                        docs_to_index.append(doc)
                    
                    # Folosim retriever-ul pentru a adăuga documentele în document store în batch - versiune compatibilă cu farm-haystack 1.26.4
                    retriever = EmbeddingRetriever(
                        document_store=document_store,
                        model_name_or_path="sentence-transformers/all-MiniLM-L6-v2",
                        device="cpu",
                        batch_size=32  # Procesare în batch pentru performanță mai bună
                    )
                    
                    # Actualizăm embedding-urile documentelor într-un singur batch
                    logger.info(f"Generare embeddings pentru {len(docs_to_index)} documente", 
                               extra={"docs_count": len(docs_to_index)})
                    document_store.update_embeddings(retriever=retriever, documents=docs_to_index)
                    
                    # Scriem toate documentele în document store într-un singur batch
                    document_store.write_documents(docs_to_index)
                    
                    logger.info(f"Au fost indexate {len(docs_to_index)} chunk-uri din JSON", 
                               extra={"indexed_count": len(docs_to_index), "file": file.filename})
                    
                    # Curățăm fișierul temporar
                    os.remove(file_path)
                    
                    return {
                        "status": "success", 
                        "message": f"Fișierul JSON a fost procesat cu succes. {len(docs_to_index)} chunk-uri au fost indexate.",
                        "chunks_count": len(docs_to_index)
                    }
                else:
                    # Dacă nu sunt chunk-uri predefinite, procesăm ca JSON normal
                    with open(file_path, "r", encoding='utf-8') as f:
                        json_data = json.load(f)
                    content = json.dumps(json_data, ensure_ascii=False, indent=2)
                    logger.info("JSON fără chunk-uri procesat ca text normal", extra={"file": file.filename})
            except json.JSONDecodeError as je:
                logger.error(f"Eroare la parsarea JSON: {str(je)}", extra={"file": file.filename, "error": str(je)})
                raise HTTPException(status_code=400, detail=f"Fișierul JSON nu este valid: {str(je)}")
            except Exception as e:
                logger.error(f"Eroare la procesarea JSON: {str(e)}", extra={"file": file.filename, "error": str(e)})
                raise HTTPException(status_code=400, detail=f"Eroare la procesarea JSON: {str(e)}")
        else:
            print(f"Tip fișier neacceptat: {file.filename}")
            raise HTTPException(status_code=400, detail="Tip de fișier neacceptat. Folosiți PDF, DOCX, TXT, JSON.")
        
        # Verificăm dacă am extras conținut valid
        if not content or len(content.strip()) < 10:
            raise HTTPException(status_code=400, detail="Nu s-a putut extrage conținut valid din fișier")
            
        print(f"Conținut extras, lungime: {len(content)} caractere")
        
        # Creăm documentul cu metadate relevante
        doc_to_process = Document(
            content=content, 
            meta={
                "source": file.filename,
                "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
                "file_size": len(content_bytes),
                "file_type": file_ext.split(".")[-1]
            }
        )

        # Inițializăm document store și configurăm procesarea documentelor
        print("Inițializare document store...")
        document_store = get_document_store(collection_name)
        
        print("Procesare document...")
        # Folosim preprocessor pentru a fragmenta documentul
        processed_docs = preprocessor.process(documents=[doc_to_process])
        
        print(f"Document fragmentat în {len(processed_docs)} bucăți")
        
        # Inițializăm retriever-ul pentru embedding - versiune compatibilă cu farm-haystack 1.26.4
        retriever = EmbeddingRetriever(
            document_store=document_store,
            model_name_or_path="sentence-transformers/all-MiniLM-L6-v2",
            device="cpu"
        )
        
        # Actualizăm embedding-urile documentelor
        document_store.update_embeddings(retriever=retriever, documents=processed_docs)
        
        # Scriem documentele în document store
        document_store.write_documents(processed_docs)
        
        processing_time = time.time() - start_time
        print(f"Procesare finalizată în {processing_time:.2f} secunde")
        
        print("Curățare fișier temporar...")
        os.remove(file_path)
        
        print("Procesare finalizată cu succes!")
        return {"status": "success", "filename": file.filename, "collection": collection_name}
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"EROARE LA PROCESARE: {str(e)}\n{error_details}")
        raise HTTPException(status_code=500, detail=f"A apărut o eroare: {str(e)}")

# Funcție reutilizabilă pentru interogarea colecției
async def query_documents(collection_name: str, query_text: str, top_k: int = 5):
    """
    Interoghează o colecție și returnează documentele relevante.
    
    Args:
        collection_name: Numele colecției de interogat
        query_text: Textul interogării
        top_k: Numărul maxim de rezultate de returnat
        
    Returns:
        Lista de documente relevante pentru interogare
    """
    try:
        # Verificăm dacă interogarea este goală
        if not query_text.strip():
            logger.warning("Interogare goală primită", extra={"collection": collection_name})
            return []
            
        # Generăm o cheie pentru cache
        cache_key = f"{collection_name}:{query_text}:{top_k}"
        
        # Verificăm cache-ul pentru a evita procesarea repetată
        cached_result = get_from_query_cache(cache_key)
        if cached_result is not None:
            logger.info(f"Rezultat găsit în cache pentru interogarea: '{query_text[:30]}...'")
            return cached_result
            
        start_time = time.time()
        with TimingLogger("Procesare interogare", collection=collection_name):
            # Inițializăm document store
            document_store = get_document_store(collection_name)
            
            # 1. Căutare exactă - verificăm dacă avem o potrivire exactă a frazei
            all_docs = document_store.get_all_documents()
            exact_matches = []
            
            # Pregătim variante ale interogării pentru căutare mai flexibilă
            query_variants = [
                query_text,  # Original
                query_text.lower(),  # Lowercase
                ' '.join(query_text.lower().split()),  # Normalizat (fără spații multiple)
                re.sub(r'[^\w\s]', '', query_text.lower())  # Fără punctuație
            ]
            
            # Căutăm potriviri exacte folosind toate variantele
            for doc in all_docs:
                content = doc.content
                content_lower = content.lower()
                content_norm = ' '.join(content_lower.split())
                content_clean = re.sub(r'[^\w\s]', '', content_lower)
                
                # Verificăm dacă oricare dintre variante se potrivește
                for variant in query_variants:
                    if variant in content or variant in content_lower or \
                       variant in content_norm or variant in content_clean:
                        exact_matches.append({
                            "content": content,
                            "score": 1.0,  # Scor maxim pentru potrivire exactă
                            "meta": doc.meta,
                            "id": doc.id,
                            "embedding": None,
                            "match_type": "exact"
                        })
                        break  # Trecem la următorul document dacă am găsit o potrivire
            
            # Dacă avem suficiente potriviri exacte, le returnăm direct
            if len(exact_matches) >= top_k:
                logger.info(f"Găsite {len(exact_matches)} potriviri exacte în {time.time() - start_time:.2f} secunde")
                return exact_matches[:top_k]
            
            # 2. Căutare semantică (embeddings)
            # Folosim retriever-ul pentru căutare semantică
            logger.info("Utilizare retriever pentru căutare semantică")
            
            # Inițializăm retriever-ul cu document store-ul - versiune compatibilă cu farm-haystack 1.26.4
            retriever = EmbeddingRetriever(
                document_store=document_store,
                model_name_or_path="sentence-transformers/all-MiniLM-L6-v2",
                top_k=top_k * 2,  # Cerem mai multe rezultate pentru a avea opțiuni
                device="cpu"
            )
            
            # În versiunea nouă nu mai este nevoie de embed_queries = True
            
            # Obținem rezultatele direct din retriever
            semantic_results = retriever.run(query=query_text)
            
            # Convertim rezultatele semantice în dicționare pentru a putea adăuga match_type
            semantic_results_dicts = []
            for res in semantic_results:
                if isinstance(res, dict):
                    res_dict = res.copy()
                    res_dict["match_type"] = "semantic"
                    semantic_results_dicts.append(res_dict)
                else:
                    # Dacă rezultatul este un Document sau alt tip de obiect
                    semantic_results_dicts.append({
                        "content": res.content if hasattr(res, "content") else str(res),
                        "score": getattr(res, "score", 0.8),  # Scor implicit pentru rezultate semantice
                        "meta": getattr(res, "meta", {}),
                        "id": getattr(res, "id", f"sem_{id(res)}"),  # Generăm un ID unic dacă nu există
                        "embedding": None,
                        "match_type": "semantic"
                    })
            
            semantic_results = semantic_results_dicts
            
            # 3. Căutare bazată pe cuvinte cheie
            keyword_results = []
            
            # Lista de cuvinte comune în limba română care pot fi ignorate
            stopwords = {'si', 'in', 'la', 'cu', 'de', 'pe', 'pentru', 'din', 'care', 'este', 'sunt', 
                        'ce', 'cum', 'unde', 'cand', 'cat', 'a', 'al', 'ai', 'ale', 'o', 'un', 'una', 
                        'niste', 'acest', 'aceasta', 'aceste', 'acesti', 'sau', 'ori', 'dar', 'insa'}
            
            # Extragem termenii semnificativi (fără stopwords)
            query_terms_raw = query_text.lower().split()
            query_terms = [term for term in query_terms_raw if term not in stopwords and len(term) > 2]
            
            # Dacă nu avem termeni semnificativi, folosim toți termenii
            if not query_terms:
                query_terms = query_terms_raw
            
            # Calculăm scoruri pentru potrivirea cuvintelor cheie
            for doc in all_docs:
                content_lower = doc.content.lower()
                
                # Numărăm potrivirile pentru fiecare termen
                matches = sum(1 for term in query_terms if term in content_lower)
                
                # Calculăm un scor mai sofisticat
                if matches > 0 and len(query_terms) > 0:
                    # Scorul de bază este proporția termenilor găsiți
                    base_score = matches / len(query_terms)
                    
                    # Bonus pentru potriviri consecutive (fraze)
                    consecutive_bonus = 0
                    for i in range(len(query_terms) - 1):
                        if f"{query_terms[i]} {query_terms[i+1]}" in content_lower:
                            consecutive_bonus += 0.1
                    
                    # Scorul final
                    final_score = min(0.95, base_score + consecutive_bonus)  # Maxim 0.95 pentru a fi sub potrivirile exacte
                    
                    keyword_results.append({
                        "content": doc.content,
                        "score": final_score,
                        "meta": doc.meta,
                        "id": doc.id,
                        "embedding": None,
                        "match_type": "keyword"
                    })
            
            # Sortăm rezultatele după scor
            keyword_results = sorted(keyword_results, key=lambda x: x["score"], reverse=True)[:top_k]
            
            # Combinăm toate rezultatele (exacte, semantice și bazate pe cuvinte cheie)
            combined_results = exact_matches + semantic_results + keyword_results
            
            # Eliminăm duplicatele (păstrăm prima apariție a fiecărui document)
            seen_ids = set()
            unique_results = []
            
            for result in combined_results:
                # Verificăm dacă result este un dicționar și are cheia 'id'
                if isinstance(result, dict) and "id" in result:
                    if result["id"] not in seen_ids:
                        seen_ids.add(result["id"])
                        unique_results.append(result)
                # Dacă este Document din Haystack
                elif hasattr(result, "id"):
                    if result.id not in seen_ids:
                        seen_ids.add(result.id)
                        # Convertim Document la dicționar pentru consistență
                        unique_results.append({
                            "content": result.content,
                            "score": getattr(result, "score", 0.0),
                            "meta": result.meta,
                            "id": result.id,
                            "embedding": None,
                            "match_type": "semantic"
                        })
            
            # Sortăm rezultatele finale după scor
            final_results = sorted(unique_results, key=lambda x: x.get("score", 0.0) if isinstance(x, dict) else 0.0, reverse=True)[:top_k]
            
            # Salvăm rezultatul în cache pentru viitoare interogări similare
            save_to_query_cache(cache_key, final_results)
            
            logger.info(f"Interogare finalizată, găsite {len(final_results)} rezultate")
            return final_results
    except Exception as e:
        logger.error(f"Eroare la interogarea colecției: {str(e)}")
        raise e

@app.post("/collections/{collection_name}/query", tags=["CRUD"])
async def query_collection(collection_name: str, request: Request):
    """
    Endpoint pentru interogarea documentelor dintr-o folder.
    
    Args:
        collection_name: Numele folder de interogat
        request: Cererea cu parametrii pentru interogare
    
    Returns:
        Lista de documente relevante pentru interogare
    """
    try:
        data = await request.json()
        query_text = data.get("question", "")
        top_k = data.get("top_k", 5)
        
        logger.info(f"Delegare interogare către funcția reutilizabilă query_documents", 
                   extra={"collection": collection_name, "query": query_text[:30] if query_text else "", "top_k": top_k})
        
        # Apelăm direct funcția query_documents reutilizabilă
        results = await query_documents(collection_name, query_text, top_k)
        return results
        
    except Exception as e:
        logger.error(f"Eroare la procesarea cererii: {str(e)}")
        raise HTTPException(status_code=500, detail=f"A apărut o eroare la procesarea interogării: {str(e)}")

# READ: Endpoint pentru a lista toate folderele
@app.get("/collections", tags=["Management"])
async def list_collections():
    collections_response = qdrant_client.get_collections()
    return [c.name for c in collections_response.collections]

# READ: Endpoint pentru a lista documentele dintr-un folder
@app.get("/collections/{collection_name}/documents", tags=["Management"])
async def list_documents(collection_name: str):
    try:
        document_store = get_document_store(collection_name)
        documents = document_store.filter_documents()
        
        # Grupăm documentele după fișierul sursă
        files_dict = {}
        
        for doc in documents:
            # Extragem metadatele relevante
            meta = doc.meta if hasattr(doc, 'meta') else {}
            source = meta.get('source', 'Necunoscut')
            created_at = meta.get('created_at', 'Necunoscut')
            
            # Dacă fișierul sursă nu există în dicționar, îl adăugăm
            if source not in files_dict:
                files_dict[source] = {
                    "source": source,
                    "created_at": created_at,
                    "doc_count": 0,
                    "doc_ids": []
                }
            
            # Incrementăm numărul de documente pentru acest fișier
            files_dict[source]["doc_count"] += 1
            files_dict[source]["doc_ids"].append(doc.id)
        
        # Convertăm dicționarul într-o listă pentru a o returna
        result = list(files_dict.values())
        
        # Sortăm rezultatul după data creării (dacă există)
        result.sort(key=lambda x: x.get("created_at", ""), reverse=True)
        
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Eroare la listarea documentelor: {str(e)}")

# CREATE: Endpoint pentru a crea un folder nouă
@app.post("/collections/{collection_name}", tags=["Management"])
async def create_collection(collection_name: str):
    """
    Endpoint pentru crearea unei foldere noi în Qdrant.
    
    Args:
        collection_name: Numele folder de creat
    
    Returns:
        Informații despre rezultatul operației
    """
    try:
        with TimingLogger("Încărcare fișier", collection=collection_name):
            # Verificăm dacă colecția există deja
            collections_response = qdrant_client.get_collections()
            existing_collections = [c.name for c in collections_response.collections]
            
            if collection_name not in existing_collections:
                logger.info(f"Colecția {collection_name} nu există. Se creează automat.", extra={"collection": collection_name})
                qdrant_client.create_collection(
                    collection_name=collection_name,
                    vectors_config={
                        "content": {
                            "size": 768,  # Dimensiunea vectorului pentru modelul all-mpnet-base-v2
                            "distance": "Cosine"
                        }
                    }
                )
                return {"status": "created", "collection": collection_name}
            else:
                return {"status": "exists", "message": f"Colecția {collection_name} există deja"}
    except Exception as e:
        logger.error(f"Eroare la crearea folder: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Eroare la crearea folder: {str(e)}")

# DELETE: Endpoint pentru a șterge documentele asociate cu un fișier
@app.delete("/collections/{collection_name}/documents", tags=["Management"])
async def delete_documents_by_source(collection_name: str, request: Request):
    data = await request.json()
    source_file = data.get("source")
    if not source_file:
        raise HTTPException(status_code=400, detail="Trebuie specificat un 'source'.")
    
    try:
        # Obținem document store
        document_store = get_document_store(collection_name)
        
        # Mai întâi trebuie să găsim toate documentele cu sursa specificată
        documents = document_store.filter_documents()
        
        # Filtrăm documentele care au sursa specificată
        docs_to_delete = [doc.id for doc in documents if doc.meta.get("source") == source_file]
        
        if not docs_to_delete:
            return {"status": "no_documents_found", "source": source_file, "collection": collection_name}
        
        # Ștergem documentele folosind ID-urile lor
        document_store.delete_documents(document_ids=docs_to_delete)
        
        return {"status": "deleted", "source": source_file, "collection": collection_name, "count": len(docs_to_delete)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Eroare la ștergerea documentelor: {str(e)}")

# DELETE: Endpoint pentru a șterge o întreagă folder
@app.delete("/collections/{collection_name}", tags=["Management"])
async def delete_collection(collection_name: str):
    try:
        # Apelul nu este asincron, nu trebuie să folosim await
        result = qdrant_client.delete_collection(collection_name=collection_name)
        return {"status": "collection_deleted", "collection": collection_name, "result": result}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Eroare la ștergerea folder: {str(e)}")

# Endpoint pentru generarea răspunsurilor folosind Gemini

# Endpoint pentru generarea răspunsurilor folosind Gemini
@app.post("/collections/{collection_name}/generate", tags=["Generation"])
async def generate_response(collection_name: str, request: GenerateRequest):
    """
    Endpoint pentru generarea răspunsurilor folosind Gemini pe baza documentelor recuperate.
    
    Args:
        collection_name: Numele folder
        request: Parametrii pentru generare
    
    Returns:
        Răspunsul generat de Gemini
    """
    try:
        with TimingLogger("Generare răspuns Gemini", collection=collection_name, query=request.query[:50]):
            # Extragem parametrii din cerere
            query = request.query
            temperature = request.temperature if request.temperature is not None else 0.2
            top_k_docs = request.top_k_docs if request.top_k_docs is not None else 5
            
            logger.info(f"Parametri primiti pentru generare", 
                       extra={"query": query[:50], "temperature": temperature, "top_k_docs": top_k_docs})
            
            # Verificăm dacă interogarea este goală
            if not query.strip():
                logger.warning("Interogare goală primită", extra={"collection": collection_name})
                return {"answer": "Interogarea nu poate fi goală."}
            
            # Pas 1: Interogăm colecția pentru a obține documentele relevante folosind funcția reutilizabilă
            logger.info(f"Recuperare documente pentru generare folosind funcția reutilizabilă query_documents", 
                       extra={"collection": collection_name, "query": query[:50], "top_k": top_k_docs})
            
            # Utilizăm direct funcția query_documents în loc să simulăm un apel HTTP intern
            docs_response = await query_documents(collection_name, query, top_k_docs)
            
            # Verificăm dacă am găsit documente
            if not docs_response or len(docs_response) == 0:
                logger.warning("Nu s-au găsit documente pentru generarea răspunsului", 
                             extra={"collection": collection_name, "query": query[:50]})
                return {"answer": "Nu am găsit informații relevante pentru a răspunde la această întrebare.", 
                        "documents": []}
            
            # Pas 2: Folosim Gemini pentru a genera un răspuns bazat pe documentele recuperate
            logger.info(f"Generare răspuns cu Gemini pentru {len(docs_response)} documente", 
                       extra={"collection": collection_name, "doc_count": len(docs_response)})
            
            gemini = GeminiGenerator()
            answer = gemini.generate_response(
                query=query,
                context_docs=docs_response,
                temperature=temperature
            )
            
            return {"answer": answer, "documents": docs_response}
    except Exception as e:
        logger.error(f"Eroare la generarea răspunsului", 
                    extra={"error": str(e), "collection": collection_name, "traceback": traceback.format_exc()})
        return {"error": f"A apărut o eroare la generarea răspunsului: {str(e)}", "answer": ""}

# --- Pornire Uvicorn (pentru rulare locală) ---
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("rag_api:app", host="0.0.0.0", port=8070, reload=True)